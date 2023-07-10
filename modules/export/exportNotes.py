import PyQt5.QtWidgets as qt
from PyQt5.QtGui import QIcon
from PyQt5.QtCore import QSize, QThread
from PyQt5 import uic
import os
import docx
import webbrowser
import sqlite3
import json
import shutil
from distutils.dir_util import copy_tree
import zipfile
import re
import string
from PIL import Image
from modules.modImage import getCentralCoordinates

window = None
POWER_SYMBOL = {"0": "⁰", "1": "¹", "2": "²", "3": "³", "4": "⁴", "5": "⁵", "6": "⁶", "7": "⁷", "8": "⁸", "9": "⁹",
                "-": "⁻"}


class GenerateDocThread(QThread):
    """ Generates word document """

    def run(self):
        global window
        doc = Doc()  # Create doc object
        doc.addTitle(f"{window.courseName} Notes")
        conn = sqlite3.connect(window.databasePath)
        c = conn.cursor()

        c.execute("SELECT topicID, topicName FROM topics ORDER BY topicID")
        topics = c.fetchall()

        for topic in topics:
            doc.addHeading(topic[1])
            c.execute(f"SELECT subtopicID, subtopicName FROM subtopics WHERE topicID = {topic[0]} "
                      "ORDER BY subtopicID")
            subtopics = c.fetchall()

            for subtopic in subtopics:
                doc.addSubheading(subtopic[1])
                c.execute(f"SELECT type, question, answer, noteID, qImageNo, aImageNo FROM notes WHERE \
                            subtopicID = {subtopic[0]} ORDER BY noteID")
                notes = c.fetchall()

                for note in notes:
                    question = note[1].replace("`", "'")
                    answer = note[2].replace("`", "'")

                    # Question diagram
                    if note[4] == 0:
                        qImageNo = None
                    else:
                        qImageNo = f"data/{window.courseName}/images/{note[4]}.png"

                    # Answer diagram
                    if note[5] == 0:
                        aImageNo = None
                    else:
                        aImageNo = f"data/{window.courseName}/images/{note[5]}.png"

                    # Clean powers
                    powers = re.findall(r"\^\-?\d*", question)
                    for power in powers:
                        newPower = "".join([POWER_SYMBOL[i] for i in power[1:]])
                        question = question.replace(power, newPower)

                    powers = re.findall(r"\^\-?\d*", answer)
                    for power in powers:
                        newPower = "".join([POWER_SYMBOL[i] for i in power[1:]])
                        answer = answer.replace(power, newPower)

                    # Split up string
                    answerList = answer.split("<br><br>")
                    answerList = [line.replace("<br>", "") for line in answerList]

                    if note[0] == "Fact":
                        doc.addFact(question, answerList, qImageNo)

                    elif note[0] == "Definition":
                        doc.addDef(question, answerList, qImageNo)

                    elif note[0] == "Formula":
                        c.execute(f"SELECT name, symbol, unit, min, max, step FROM formulas f JOIN notes n ON \
                                            f.formulaID = n.formulaID WHERE noteID = {note[3]}")
                        formulaVariables = [
                            {"name": term[0], "symbol": term[1], "unit": term[2], "min": term[3], "max": term[4],
                             "step": term[5]} for term in c.fetchall()]

                        # Reorganise formulaVariables
                        text = note[2]
                        for index, char in enumerate(note[2]):
                            if char.lower() not in list(string.ascii_lowercase) + list(string.digits) + ["_", "="]:
                                text = text[:index] + " " + text[index + 1:]
                        text = text.replace("=", " = ")
                        terms = [var for var in text.split() if not var.isnumeric()]
                        terms = list(dict.fromkeys(terms))
                        sep = terms.index("=")
                        lhs, rhs = terms[:sep], terms[sep + 1:]
                        formulaVariables = [[term for term in formulaVariables if term["symbol"] == var][0] for var
                                            in lhs + rhs]

                        doc.addFormula(answer, formulaVariables, lhs, rhs)

                    elif note[0] == "Process":
                        c.execute(f"SELECT processID FROM notes WHERE noteID = {note[3]}")
                        processID = c.fetchone()[0]
                        c.execute(f"SELECT step FROM processes WHERE processID = {processID} ORDER BY stepNo")
                        answerList = [f"{i}. {line}" for i, line in enumerate([step[0] for step in c.fetchall()],
                                                                              start=1)]
                        doc.addProcess(question, answerList, qImageNo)

                    elif note[0] == "Diagram":
                        doc.addDiagram(question, answerList, aImageNo)

                    elif note[0] == "Table":
                        c.execute(f"SELECT tableID FROM notes WHERE noteID = {note[3]}")
                        tableID = c.fetchone()[0]
                        c.execute(
                            f"SELECT rowNo, colNo, textVal, userFill FROM tableElements WHERE tableID = {tableID} "
                            "ORDER BY rowNo, colNo")
                        values = c.fetchall()
                        noRows, noCols = values[-1][0] + 1, values[-1][1] + 1

                        tableValues = [["" for _ in range(noCols)] for _ in range(noRows)]

                        for cell in values:
                            a, b = cell[0], cell[1]
                            tableValues[a][b] = cell[2].replace("`", "'")

                        doc.addTable(question, tableValues)

        # Delete unnecessary files
        if os.path.exists("images/temp/foo.png"):
            os.remove("images/temp/foo.png")

        # Save Document
        doc.save(window.filename)
        webbrowser.open(window.filename)
        conn.close()
        window.setWindowTitle("Smart Retain")
        window.wordBtn.setText("File Generated ")
        window.wordBtn.setEnabled(True)


class Doc:
    def __init__(self):
        self.doc = docx.Document()
        self.paragraphNo = 0
        self.beforeTitle = False

        sections = self.doc.sections
        for section in sections:
            section.top_margin = docx.shared.Cm(1.27)
            section.bottom_margin = docx.shared.Cm(1.27)
            section.left_margin = docx.shared.Cm(1.27)
            section.right_margin = docx.shared.Cm(1.27)

        # Create styles
        obj_styles = self.doc.styles

        # Document Title
        obj_charstyle = obj_styles.add_style("0", docx.enum.style.WD_STYLE_TYPE.CHARACTER)
        obj_font = obj_charstyle.font
        obj_font.bold, obj_font.italic, obj_font.underline = True, False, True
        obj_font.name, obj_font.size = "Arial Black", docx.shared.Pt(36)

        # Heading
        obj_charstyle = obj_styles.add_style("1", docx.enum.style.WD_STYLE_TYPE.CHARACTER)
        obj_font = obj_charstyle.font
        obj_font.bold, obj_font.italic, obj_font.underline = True, False, True
        obj_font.name, obj_font.size = "Arial", docx.shared.Pt(24)

        # Subtitle
        obj_charstyle = obj_styles.add_style("2", docx.enum.style.WD_STYLE_TYPE.CHARACTER)
        obj_font = obj_charstyle.font
        obj_font.bold, obj_font.italic, obj_font.underline = False, True, False
        obj_font.name, obj_font.size = "Arial", docx.shared.Pt(18)

        # Fact Question
        obj_charstyle = obj_styles.add_style("3", docx.enum.style.WD_STYLE_TYPE.CHARACTER)
        obj_font = obj_charstyle.font
        obj_font.bold, obj_font.italic, obj_font.underline = True, True, False
        obj_font.name, obj_font.size = "Arial", docx.shared.Pt(12)
        obj_font.color.rgb = docx.shared.RGBColor.from_string("1F497D")

        # Definition Question
        obj_charstyle = obj_styles.add_style("4", docx.enum.style.WD_STYLE_TYPE.CHARACTER)
        obj_font = obj_charstyle.font
        obj_font.bold, obj_font.italic, obj_font.underline = True, True, False
        obj_font.name, obj_font.size = "Arial", docx.shared.Pt(12)
        obj_font.color.rgb = docx.shared.RGBColor.from_string("1CADE4")

        # Text
        obj_charstyle = obj_styles.add_style("5", docx.enum.style.WD_STYLE_TYPE.CHARACTER)
        obj_font = obj_charstyle.font
        obj_font.bold, obj_font.italic, obj_font.underline = False, False, False
        obj_font.name, obj_font.size = "Arial", docx.shared.Pt(12)

        # Bold Text
        obj_charstyle = obj_styles.add_style("6", docx.enum.style.WD_STYLE_TYPE.CHARACTER)
        obj_font = obj_charstyle.font
        obj_font.bold, obj_font.italic, obj_font.underline = True, False, False
        obj_font.name, obj_font.size = "Arial", docx.shared.Pt(12)

        # Formula Text
        obj_charstyle = obj_styles.add_style("7", docx.enum.style.WD_STYLE_TYPE.CHARACTER)
        obj_font = obj_charstyle.font
        obj_font.bold, obj_font.italic, obj_font.underline = True, True, False
        obj_font.name, obj_font.size = "Arial", docx.shared.Pt(12)
        obj_font.color.rgb = docx.shared.RGBColor.from_string("62A39F")

        # Process Text
        obj_charstyle = obj_styles.add_style("8", docx.enum.style.WD_STYLE_TYPE.CHARACTER)
        obj_font = obj_charstyle.font
        obj_font.bold, obj_font.italic, obj_font.underline = True, True, False
        obj_font.name, obj_font.size = "Arial", docx.shared.Pt(12)
        obj_font.color.rgb = docx.shared.RGBColor.from_string("2683C6")

        # Diagram Text
        obj_charstyle = obj_styles.add_style("9", docx.enum.style.WD_STYLE_TYPE.CHARACTER)
        obj_font = obj_charstyle.font
        obj_font.bold, obj_font.italic, obj_font.underline = True, True, False
        obj_font.name, obj_font.size = "Arial", docx.shared.Pt(12)
        obj_font.color.rgb = docx.shared.RGBColor.from_string("624CB4")

        # Table Text
        obj_charstyle = obj_styles.add_style("10", docx.enum.style.WD_STYLE_TYPE.CHARACTER)
        obj_font = obj_charstyle.font
        obj_font.bold, obj_font.italic, obj_font.underline = True, True, False
        obj_font.name, obj_font.size = "Arial", docx.shared.Pt(12)
        obj_font.color.rgb = docx.shared.RGBColor.from_string("3E8853")

    def addTitle(self, text):
        """Adds paragraph to document"""
        paragraph = self.doc.add_paragraph("")
        paragraph.add_run(text, style="0")
        self.beforeTitle = True

    def addHeading(self, text):
        """Adds paragraph to document"""
        if not self.beforeTitle:
            self.doc.add_page_break()
        paragraph = self.doc.add_paragraph("")
        paragraph.add_run(text, style="1")
        self.beforeTitle = False

    def addSubheading(self, text):
        """Adds paragraph to document"""
        paragraph = self.doc.add_paragraph("")
        paragraph.add_run(text, style="2")
        
    def addImage(self, path):
        """ Adds image to document """
        img = Image.open(path)
        WIDTH = 720
        if img.width == WIDTH:
            # Standard image
            t = 20

            # Find coordinates of main image
            n, s, w, e = getCentralCoordinates(img, t)
            img = img.crop((w, n, e, s))

        elif img.width > WIDTH:
            scale = WIDTH / img.width
            img = img.resize((int(img.width * scale), int(img.height * scale)))

        img.save("images/temp/foo.png")

        p = self.doc.add_paragraph("")
        r = p.add_run()
        r.add_picture("images/temp/foo.png", width=docx.shared.Inches(img.width * 0.008),
                      height=docx.shared.Inches(img.height * 0.008))

    def addFact(self, questionText, answerText, qImageNo):
        """Adds paragraph to document"""
        paragraph = self.doc.add_paragraph("")
        paragraph.add_run(questionText + ": ", style="3")
        
        # If there is an image
        if qImageNo:
            answerText.insert(0, "")
            self.addImage(qImageNo)
        
        for i, text in enumerate(answerText):
            for word in text.split():

                if word[0] == "*" and word[-1] == "*" and len(word) > 1:
                    _style = "6"
                    word = word[1:-1]
                else:
                    _style = "5"

                paragraph.add_run(word + " ", style=_style)

            if i < len(answerText) - 1:
                paragraph = self.doc.add_paragraph("")

    def addDef(self, questionText, answerText, qImageNo):
        """Adds paragraph to document"""
        paragraph = self.doc.add_paragraph("")
        paragraph.add_run(questionText + ": ", style="4")

        # If there is an image
        if qImageNo:
            answerText.insert(0, "")
            self.addImage(qImageNo)
        
        for i, text in enumerate(answerText):
            for word in text.split():

                if word[0] == "*" and word[-1] == "*" and len(word) > 1:
                    _style = "6"
                    word = word.replace("*", "")
                else:
                    _style = "5"
                paragraph.add_run(word + " ", style=_style)

            if i < len(answerText) - 1:
                paragraph = self.doc.add_paragraph("")

    def addProcess(self, questionText, answerText, qImageNo):
        """Adds paragraph to document"""
        paragraph = self.doc.add_paragraph("")
        paragraph.add_run(questionText + ": ", style="8")
        
        # If there is an image
        if qImageNo:
            answerText.insert(0, "")
            self.addImage(qImageNo)
        else:
            paragraph = self.doc.add_paragraph("")
            
        for i, text in enumerate(answerText):
            for word in text.split():

                if word[0] == "*" and word[-1] == "*" and len(word) > 1:
                    _style = "6"
                    word = word.replace("*", "")
                else:
                    _style = "5"
                paragraph.add_run(word + " ", style=_style)

            if i < len(answerText) - 1:
                paragraph.add_run(" "*200, style="5")

    def addTable(self, questionText, tableData):
        """ Adds table question to document """
        paragraph = self.doc.add_paragraph("")
        paragraph.add_run(questionText, style="10")

        # Creating a table object
        table = self.doc.add_table(rows=0, cols=len(tableData[0]))
        table.style = "Table Grid"

        # Adding data from the list to the table
        for j, record in enumerate(tableData):
            # Adding a row and then adding data in it.
            row = table.add_row().cells

            for i, cell in enumerate(record):
                row[i].text = str(cell)
                table.cell(j, i).vertical_alignment = docx.enum.table.WD_ALIGN_VERTICAL.CENTER

        # Change font
        for row in table.rows:
            for cell in row.cells:
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    for run in paragraph.runs:
                        font = run.font
                        font.name, font.size = "Arial", docx.shared.Pt(12)

        # Change font and style for the first row
        for i, cell in enumerate(table.rows[0].cells):
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.name, run.font.size = "Arial", docx.shared.Pt(12)

        for row in table.rows:
            row.height = docx.shared.Cm(0.7)

        paragraph = self.doc.add_paragraph("")

    def addDiagram(self, questionText, answerText, aImageNo):
        """ Adds diagram question to document """
        paragraph = self.doc.add_paragraph("")
        paragraph.add_run(questionText + ": ", style="9")
        self.addImage(aImageNo)

    def addFormula(self, answerText, variables, lhs, rhs):
        """Adds paragraph to document"""
        if len(lhs) == 1:
            # Left hand side only consists of one term
            questionText = "The Formula for " + \
                           [term for term in variables if term["symbol"] == lhs[0]][0]["name"].title()
        elif len(rhs) == 1:
            # Right hand side only consists of one term
            questionText = "The Formula for " + \
                           [term for term in variables if term["symbol"] == rhs[0]][0]["name"].title()
        else:
            # Equation
            questionText = "The Equation consisting of " + ", ".join([term["name"].title() for term in
                                                                      variables[:-1]]) + " and " + variables[-1][
                               "name"].title()

        paragraph = self.doc.add_paragraph("")
        paragraph.add_run(questionText + ": ", style="7")

        # Clean text
        answerText = answerText.replace(" * ", "")

        # Clean powers
        powers = re.findall(r"\^\-?\d*", answerText)
        for power in powers:
            newPower = "".join([POWER_SYMBOL[i] for i in power[1:]])
            answerText = answerText.replace(power, newPower)

        paragraph.add_run(answerText, style="5")
        paragraph = self.doc.add_paragraph("")

        # Enter symbol meanings
        for i, term in enumerate(variables):
            symbol, name, unit = term["symbol"], term["name"].title(), term["unit"]

            # Clean unit
            powers = re.findall(r"\^\-?\d*", unit)
            for power in powers:
                newPower = "".join([POWER_SYMBOL[i] for i in power[1:]])
                unit = unit.replace(power, newPower)

            paragraph.add_run(f"{symbol}: {name} ({unit})", style="5")

            if i < len(variables) - 1:
                paragraph.add_run(" "*200, style="5")

    def save(self, filename):
        """Exports document"""
        self.doc.save(filename)


class Window(qt.QMainWindow):
    filename = None

    def __init__(self, superClass, course, color):
        """ Main Window """
        global window
        super(Window, self).__init__()
        uic.loadUi("gui/exportNotes.ui", self)
        self.setWindowIcon(QIcon("images/logo.png"))
        window = self
        self.superClass = superClass
        self.courseName = course
        self.color = color
        self.databasePath = "data/" + course + "/courseData.db"
        self.updateTitleColor()
        self.show()

        # Get widgets
        self.importBtn = self.findChild(qt.QPushButton, "importBtn")
        self.wordBtn = self.findChild(qt.QPushButton, "wordBtn")
        self.dataBtn = self.findChild(qt.QPushButton, "dataBtn")

        # Add Icons
        self.importBtn.setIcon(QIcon("images/import.png"))
        self.importBtn.setIconSize(QSize(50, 50))
        self.wordBtn.setIcon(QIcon("images/doc.png"))
        self.wordBtn.setIconSize(QSize(50, 50))
        self.dataBtn.setIcon(QIcon("images/zip.png"))
        self.dataBtn.setIconSize(QSize(50, 50))

        # Bind widgets
        self.importBtn.clicked.connect(self.showImportOptions)
        self.wordBtn.clicked.connect(self.convertToDoc)
        self.dataBtn.clicked.connect(self.exportNotesToZip)

        # Check for notes
        conn = sqlite3.connect(self.databasePath)
        c = conn.cursor()
        c.execute("SELECT name FROM sqlite_master WHERE type='table' AND name='notes'")
        notes = c.fetchall()
        conn.commit()
        conn.close()

        # Update buttons
        self.wordBtn.setEnabled(bool(notes))
        self.dataBtn.setEnabled(bool(notes))

        self.generateDoc = GenerateDocThread()

    def importNotes(self, fName, merge):
        """ Imports notes """
        unzipPath = os.path.dirname(fName) + "/unzip"
        dataFilepath = f"data/{self.courseName}"

        # Delete folder if exists
        if os.path.exists(unzipPath) and os.path.isdir(unzipPath):
            shutil.rmtree(unzipPath)

        # Unzip folder
        with zipfile.ZipFile(fName, "r") as zip_ref:
            zip_ref.extractall(unzipPath)

        if "images" in os.listdir(unzipPath) and "courseData.db" in os.listdir(unzipPath):
            # Correctly formatted zip

            # Modify database file
            conn = sqlite3.connect(f"{unzipPath}/courseData.db")
            c = conn.cursor()
            c.execute("UPDATE notes SET practiceCount = 0, correctCount = 0, userScore = 0, starred = 0")
            conn.commit()
            conn.close()

            if not merge:
                # Delete current files and replace with new files
                os.remove(f"{dataFilepath}/courseData.db")
                shutil.rmtree(f"{dataFilepath}/images")

                # Move files
                shutil.move(f"{unzipPath}/courseData.db", f"{dataFilepath}/courseData.db")
                shutil.move(f"{unzipPath}/images", f"{dataFilepath}")

            else:
                # Get number of current database images
                conn = sqlite3.connect(self.databasePath)
                c = conn.cursor()
                folder = "data/" + self.courseName + "/images"
                images = sorted([int(os.path.splitext(file)[0]) for file in os.listdir(folder)])
                if images:
                    imageNo = images[-1]
                else:
                    imageNo = 0

                c.execute("SELECT topicID FROM topics ORDER BY topicID DESC")
                topicNo = c.fetchone()[0]
                c.execute("SELECT subtopicID FROM subtopics ORDER BY subtopicID DESC")
                subtopicsNo = c.fetchone()[0]
                c.execute("SELECT noteID FROM notes ORDER BY noteID DESC")
                notesNo = c.fetchone()[0]
                c.execute("SELECT processID FROM notes ORDER BY processID DESC")
                processNo = c.fetchone()[0]
                c.execute("SELECT formulaID FROM notes ORDER BY formulaID DESC")
                formulaNo = c.fetchone()[0]
                c.execute("SELECT tableID FROM notes ORDER BY tableID DESC")
                tableNo = c.fetchone()[0]
                conn.close()

                # Get number of new images
                conn = sqlite3.connect(f"{unzipPath}/courseData.db")
                c = conn.cursor()

                # Increment images in database
                c.execute(f"UPDATE notes SET qImageNo = qImageNo + {imageNo} WHERE qImageNo > 0")
                c.execute(f"UPDATE notes SET aImageNo = aImageNo + {imageNo} WHERE aImageNo > 0")
                c.execute(f"UPDATE notes SET processID = processID + {processNo} WHERE processID > 0")
                c.execute(f"UPDATE processes SET processID = processID + {processNo} WHERE processID > 0")
                c.execute(f"UPDATE notes SET tableID = tableID + {tableNo} WHERE tableID > 0")
                c.execute(f"UPDATE tableElements SET tableID = tableID + {tableNo} WHERE tableID > 0")
                c.execute(f"UPDATE notes SET formulaID = formulaID + {formulaNo} WHERE formulaID > 0")
                c.execute(f"UPDATE formulas SET formulaID = formulaID + {formulaNo} WHERE formulaID > 0")
                c.execute(f"UPDATE topics SET topicID = topicID + {topicNo}")
                c.execute(f"UPDATE subtopics SET topicID = topicID + {topicNo}")
                c.execute(f"UPDATE subtopics SET subtopicID = subtopicID + {subtopicsNo}")
                c.execute(f"UPDATE notes SET subtopicID = subtopicID + {subtopicsNo}")
                c.execute(f"UPDATE notes SET noteID = noteID + {notesNo}")
                conn.commit()

                # Increment image in folder
                n = len(os.listdir(f"{unzipPath}/images"))
                for i in reversed(range(1, n + 1)):
                    os.rename(f"{unzipPath}/images/{i}.png", f"{unzipPath}/images/{i + imageNo}.png")

                # Merge notes
                for file in os.listdir(f"{unzipPath}/images"):
                    shutil.move(f"{unzipPath}/images/{file}", f"{dataFilepath}/images/{file}")

                # Merge databases
                c.execute("SELECT * FROM topics")
                topicTable = c.fetchall()
                c.execute("SELECT * FROM subtopics")
                subtopicTable = c.fetchall()
                c.execute("SELECT * FROM notes")
                noteTable = c.fetchall()
                c.execute("SELECT * FROM formulas")
                formulaTable = c.fetchall()
                c.execute("SELECT * FROM processes")
                processTable = c.fetchall()
                c.execute("SELECT * FROM tableElements")
                tableElementsTable = c.fetchall()

                conn = sqlite3.connect(self.databasePath)
                c = conn.cursor()

                for topic in topicTable:
                    wildCard = ",".join(["?" for val in topic])
                    c.execute(f"INSERT INTO topics VALUES ({wildCard})", topic)

                for subtopic in subtopicTable:
                    wildCard = ",".join(["?" for val in subtopic])
                    c.execute(f"INSERT INTO subtopics VALUES ({wildCard})", subtopic)

                for note in noteTable:
                    wildCard = ",".join(["?" for val in note])
                    c.execute(f"INSERT INTO notes VALUES ({wildCard})", note)

                for formula in formulaTable:
                    wildCard = ",".join(["?" for val in formula])
                    c.execute(f"INSERT INTO formulas VALUES ({wildCard})", formula)

                for process in processTable:
                    wildCard = ",".join(["?" for val in process])
                    c.execute(f"INSERT INTO processes VALUES ({wildCard})", process)

                for element in tableElementsTable:
                    wildCard = ",".join(["?" for val in element])
                    c.execute(f"INSERT INTO tableElements VALUES ({wildCard})", element)

                conn.commit()

                conn.close()

            qt.QMessageBox.information(self, "Success", "Files have been Successfully Imported", qt.QMessageBox.Ok)

        else:
            # Incorrect file uploaded
            qt.QMessageBox.critical(self, "Error", "Incorrectly formatted file", qt.QMessageBox.Ok)

        shutil.rmtree(unzipPath)

    def showImportOptions(self):
        """ Allows user to import notes """
        with open("text/currentSettings.json") as f:
            data = json.load(f)

        fName, _ = qt.QFileDialog.getOpenFileName(self, "Select Notes Zip", data["textbookPath"], "Zip files (*.zip)")

        if fName:
            # Ask whether they want to replace notes or merge notes
            conn = sqlite3.connect(self.databasePath)
            c = conn.cursor()
            c.execute("SELECT name FROM sqlite_master WHERE type='table' AND name='subtopics'")
            if c.fetchall():
                reply = qt.QMessageBox.question(self, "Import Options",
                                                "Would you like to Merge new Notes with existing notes",
                                                qt.QMessageBox.Yes | qt.QMessageBox.No | qt.QMessageBox.Cancel)
            else:
                reply = qt.QMessageBox.No
            conn.close()

            if reply == qt.QMessageBox.Yes:
                # Merge notes
                self.importNotes(fName, merge=True)

            elif reply == qt.QMessageBox.No:
                check = qt.QMessageBox.warning(self, "Are you sure?", "All existing Notes will be lost",
                                               qt.QMessageBox.Yes | qt.QMessageBox.Cancel)

                if check == qt.QMessageBox.Yes:
                    # Delete notes and replace with new ones
                    self.importNotes(fName, merge=False)

    def exportNotesToZip(self):
        """ Sends notes folder to zip """
        with open("text/currentSettings.json") as f:
            data = json.load(f)

        filename = str(qt.QFileDialog.getExistingDirectory(self, "Select where to save Notes", data["textbookPath"]))

        if filename:
            # Duplicate folder
            directoryName = f"images/temp/{self.courseName}"
            copy_tree(f"data/{self.courseName}", f"images/temp/{self.courseName}")
            os.remove(f"images/temp/{self.courseName}/courseInfo.json")

            # Save zip
            zipName = f"{filename}/{self.courseName}"
            shutil.make_archive(zipName, "zip", directoryName)
            shutil.rmtree(directoryName)

            qt.QMessageBox.information(self, "Export Successful", f"{self.courseName}.zip saved in {filename}",
                                       qt.QMessageBox.Ok)

            webbrowser.open(zipName + ".zip")

    def convertToDoc(self):
        """ Inserts values from database into document """
        with open("text/currentSettings.json") as f:
            data = json.load(f)

        path = f"{data['textbookPath']}/{self.courseName} Notes.docx"
        self.filename, _ = qt.QFileDialog.getSaveFileName(self, "Select where to save Document", path,
                                                     "Word Document (*.docx);;Word Document (*.doc)")

        if self.filename:
            self.generateDoc.start()
            self.setWindowTitle("Generating Word Document - Please wait")
            self.wordBtn.setText("Generating... Please Wait ")
            self.wordBtn.setEnabled(False)

    def closeEvent(self, event):
        """ Run when window gets closed """
        self.superClass.show()

    def updateTitleColor(self):
        """ Updates title with course colour """
        titleLabel = self.findChild(qt.QLabel, "titleLabel")

        fontColor = self.color.lstrip("#")
        lv = len(fontColor)
        r, g, b = tuple(int(fontColor[i:i + lv // 3], 16) for i in range(0, lv, lv // 3))
        if (r * 0.299 + g * 0.587 + b * 0.114) > 186:
            fontColor = "#000000"
        else:
            fontColor = "#FFFFFF"

        titleLabel.setStyleSheet("border-style: outset;border-width: 2px;border-radius: 10px;border-color: "
                                 f"#303545;background-color:{self.color}; color:{fontColor}")
